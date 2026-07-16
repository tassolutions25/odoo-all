import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    """Sets the background color of a table cell."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets the padding of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_with_spacing(doc, text, level, before=12, after=6):
    """Adds a heading with defined spacing before and after."""
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(before)
    h.paragraph_format.space_after = Pt(after)
    h.paragraph_format.keep_with_next = True
    
    # Custom heading colors
    run = h.runs[0]
    run.font.name = 'Segoe UI'
    if level == 1:
        run.font.color.rgb = RGBColor(0x4A, 0x15, 0x4B)  # Plum/Purple theme
        run.font.size = Pt(20)
        # Add a bottom border or horizontal line below Heading 1
        pPr = h._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12') # 1.5 pt
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '4A154B')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)  # Slate blue
        run.font.size = Pt(15)
    elif level == 3:
        run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)  # Muted grey
        run.font.size = Pt(12)
    return h

def main():
    doc = Document()
    
    # Configure page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set normal style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ================= COVER PAGE =================
    # Add spacing at top
    title_space = doc.add_paragraph()
    title_space.paragraph_format.space_before = Pt(80)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Ahadu Bank Payroll")
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x4A, 0x15, 0x4B)
    
    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Comprehensive Product Documentation & User Guide\nCustom Odoo 18 Module for Ethiopian Banking Operations")
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    sub_p.paragraph_format.space_after = Pt(200)

    # Metadata
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run("Prepared by: Ahadu Dev Team\nWebsite: www.ahadubank.com\nDate: June 2026\nVersion: 18.0.1.0.0")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)

    # Page Break after Cover Page
    doc.add_page_break()

    # ================= 1. OVERVIEW =================
    add_heading_with_spacing(doc, "1. Overview", level=1, before=18)
    
    p = doc.add_paragraph(
        "The Ahadu Bank Payroll (ahadu_payroll) Odoo 18 module is built on top of Odoo 18's core HR modules "
        "and the OCA base payroll engine. It addresses the unique regulatory, fiscal, and operational requirements "
        "of the Ethiopian banking sector. It provides deep integrations between employee records, work contracts, "
        "attendance data, loan schemes, tax structures, and general ledger journal distributions."
    )
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15

    p = doc.add_paragraph("Key objectives of the module include:")
    p.paragraph_format.space_after = Pt(4)
    
    bullets = [
        "Full automation of Ethiopian personal income tax (PIT) calculations.",
        "Support for local pension contributions (Employee 7%, Employer 11%).",
        "Complete automation of employee loans, approval workflows, and amortization schedules.",
        "Streamlined management of bank-specific allowances (e.g., fuel allowance, cash indemnity).",
        "Backpay calculations for retroactive salary adjustments.",
        "Integrated resignation and termination settlement processing."
    ]
    for b in bullets:
        bp = doc.add_paragraph(b, style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)

    # ================= 2. CORE FEATURES & USER GUIDE =================
    add_heading_with_spacing(doc, "2. Core Features & User Guide", level=1, before=24)

    # 2.1 Payroll Dashboard
    add_heading_with_spacing(doc, "2.1 Payroll Dashboard", level=2, before=14)
    p = doc.add_paragraph(
        "The Payroll Dashboard provides HR Managers and Finance Executives with real-time financial KPIs "
        "and visual charts illustrating payroll trends, employee distribution, active loans, and overall financial liability."
    )
    p.paragraph_format.space_after = Pt(8)

    # Insert Dashboard Image
    img_path = 'doc/images/dashboard_mockup.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6.0))
        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_p.add_run("Figure 2.1: Ahadu Bank Payroll dashboard with real-time analytics KPIs.")
        caption_run.font.size = Pt(9.5)
        caption_run.font.italic = True
        caption_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        caption_p.paragraph_format.space_before = Pt(6)
        caption_p.paragraph_format.space_after = Pt(12)

    # 2.2 Taxation & Parameter Configuration
    add_heading_with_spacing(doc, "2.2 Taxation & Parameter Configuration", level=2, before=18)
    p = doc.add_paragraph(
        "Ethiopian tax laws require progressive income brackets and specific handling of taxable allowances. "
        "The taxation engine is fully configurable via the Tax Brackets screen, featuring authorization "
        "workflows (Draft -> Submitted -> Approved) to prevent unauthorized changes."
    )
    p.paragraph_format.space_after = Pt(8)

    # Insert Tax Configuration Image
    img_path = 'doc/images/tax_config_mockup.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6.0))
        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_p.add_run("Figure 2.2: Ethiopian progressive tax bracket configuration settings.")
        caption_run.font.size = Pt(9.5)
        caption_run.font.italic = True
        caption_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        caption_p.paragraph_format.space_before = Pt(6)
        caption_p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph("Below are the standard Ethiopian progressive personal income tax brackets implemented in the module:")
    p.paragraph_format.space_after = Pt(6)

    # Create Table
    table_data = [
        ("Income Range (ETB)", "Tax Rate (%)", "Deduction (ETB)"),
        ("0 - 600", "0%", "0.00"),
        ("601 - 1,650", "10%", "60.00"),
        ("1,651 - 3,200", "15%", "142.50"),
        ("3,201 - 5,250", "20%", "302.50"),
        ("5,251 - 7,800", "25%", "565.00"),
        ("7,801 - 10,900", "30%", "955.00"),
        ("Above 10,900", "35%", "1,500.00")
    ]
    
    table = doc.add_table(rows=len(table_data), cols=3)
    table.style = 'Light Shading Accent 1'
    
    for row_idx, row_content in enumerate(table_data):
        row = table.rows[row_idx]
        for col_idx, text in enumerate(row_content):
            cell = row.cells[col_idx]
            cell.text = text
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            if row_idx == 0:
                set_cell_background(cell, "4A154B")
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif row_idx % 2 == 0:
                set_cell_background(cell, "F9F5F9")
                
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 2.3 Loan Management System
    add_heading_with_spacing(doc, "2.3 Loan Management System", level=2, before=18)
    p = doc.add_paragraph(
        "The module provides an automated employee loan management process. Employees or HR Officers can submit loan "
        "requests, structure installments, and track repayments. When a payslip is computed, the system automatically "
        "deducts the current installment due."
    )
    p.paragraph_format.space_after = Pt(8)

    # Insert Loan Management Image
    img_path = 'doc/images/loan_management_mockup.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6.0))
        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_p.add_run("Figure 2.3: Employee loan request form showing the approval workflow states.")
        caption_run.font.size = Pt(9.5)
        caption_run.font.italic = True
        caption_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        caption_p.paragraph_format.space_before = Pt(6)
        caption_p.paragraph_format.space_after = Pt(12)

    # 2.4 Overtime & Attendance Integration
    add_heading_with_spacing(doc, "2.4 Overtime & Attendance Integration", level=2, before=18)
    p = doc.add_paragraph(
        "Overtime requests are integrated with employee attendance verification. The module handles four "
        "overtime types applying multiplier rates relative to the hourly basic rate:"
    )
    p.paragraph_format.space_after = Pt(6)

    ot_bullets = [
        "Normal Overtime: Applied for hours worked after standard shifts (Multiplier: 1.25x).",
        "Night Overtime: Applied for night shifts worked between 10:00 PM and 6:00 AM (Multiplier: 1.50x).",
        "Sunday Overtime: Applied for weekend rest days (Multiplier: 2.00x).",
        "Holiday Overtime: Applied for work performed on official public holidays (Multiplier: 2.50x)."
    ]
    for ot_b in ot_bullets:
        bp = doc.add_paragraph(ot_b, style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)

    # Insert Overtime Image
    img_path = 'doc/images/overtime_attendance_mockup.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6.0))
        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_p.add_run("Figure 2.4: Overtime requests list view with rates and approval status.")
        caption_run.font.size = Pt(9.5)
        caption_run.font.italic = True
        caption_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        caption_p.paragraph_format.space_before = Pt(6)
        caption_p.paragraph_format.space_after = Pt(12)

    # 2.5 Backpay & Retroactive Adjustments
    add_heading_with_spacing(doc, "2.5 Backpay & Retroactive Adjustments", level=2, before=18)
    p = doc.add_paragraph(
        "In the event of delayed promotions or salary increments, the Backpay engine calculates the variance "
        "between what was paid and what should have been paid across a defined range of historical months. The system "
        "then automatically injects these adjusting lines into the active payroll cycle."
    )
    p.paragraph_format.space_after = Pt(8)

    # 2.6 Bonus Management
    add_heading_with_spacing(doc, "2.6 Bonus Management", level=2, before=14)
    p = doc.add_paragraph(
        "Supports performance-based and festival-based bonus payouts. Enables customized configuration of tax-exempt "
        "portions according to company policy or governing laws."
    )
    p.paragraph_format.space_after = Pt(8)

    # 2.7 Resignation & Termination Payslips
    add_heading_with_spacing(doc, "2.7 Resignation & Termination Payslips", level=2, before=14)
    p = doc.add_paragraph(
        "Automated final settlements handle employee departures smoothly. Resignation workflows calculate notice "
        "period penalties and leave encashment, while Termination workflows calculate severance payouts based on "
        "the Ethiopian Labour Proclamation (relying on years of service)."
    )
    p.paragraph_format.space_after = Pt(12)

    # ================= 3. TECHNICAL ARCHITECTURE & DATA MODELS =================
    add_heading_with_spacing(doc, "3. Technical Architecture & Data Models", level=1, before=24)
    p = doc.add_paragraph("The custom tables/models introduced by the module include:")
    p.paragraph_format.space_after = Pt(6)

    models_data = [
        ("Model Name", "Description"),
        ("hr.loan", "Employee loan request records, principal amount, amortization plan."),
        ("hr.loan.line", "Amortization installments mapped to active contract payroll lines."),
        ("ahadu.payroll.tax.bracket", "Ethiopian personal income tax brackets details."),
        ("ahadu.payroll.tax.config", "Main parameters linking tax brackets and fuel parameters."),
        ("ahadu.overtime", "Employee overtime logs, hours, rates, and approval states."),
        ("ahadu.backpay", "Retroactive backpay computations and adjustments."),
        ("cash.indemnity", "Cash risk handling allowances for tells/cashiers.")
    ]
    
    table_models = doc.add_table(rows=len(models_data), cols=2)
    table_models.style = 'Light Shading Accent 1'
    
    for row_idx, row_content in enumerate(models_data):
        row = table_models.rows[row_idx]
        for col_idx, text in enumerate(row_content):
            cell = row.cells[col_idx]
            cell.text = text
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            if row_idx == 0:
                set_cell_background(cell, "4A154B")
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif row_idx % 2 == 0:
                set_cell_background(cell, "F9F5F9")

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ================= 4. CONFIGURATION & INSTALLATION =================
    add_heading_with_spacing(doc, "4. Configuration & Installation", level=1, before=24)
    
    add_heading_with_spacing(doc, "Installation Steps", level=2, before=12)
    steps = [
        "Clone or extract the ahadu_payroll module into the custom addons folder.",
        "Restart the Odoo 18 server.",
        "Activate developer mode in Odoo, go to Apps, and click 'Update Apps List'.",
        "Search for 'Ahadu Bank Payroll' and click Install.",
        "All required dependencies (hr, hr_contract, payroll, ahadu_hr_leave, ahadu_hr) will install automatically."
    ]
    for step in steps:
        sp = doc.add_paragraph(step, style='List Number')
        sp.paragraph_format.space_after = Pt(4)

    add_heading_with_spacing(doc, "Post-Installation Setup Checklist", level=2, before=14)
    checklist = [
        "Go to Payroll > Configuration > settings and configure Tax Brackets, setting status to Approved.",
        "Verify fuel prices are updated to reflect the latest tariffs.",
        "Verify Employee bank account records are correctly mapped in hr.employee.bank.account.",
        "Configure salary rules and map them to payroll accounting journal entries."
    ]
    for chk in checklist:
        bp = doc.add_paragraph(chk, style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)

    # Save Document
    doc_path = 'doc/ahadu_payroll_documentation.docx'
    doc.save(doc_path)
    print(f"Word Document successfully saved to {doc_path}")

if __name__ == '__main__':
    main()
