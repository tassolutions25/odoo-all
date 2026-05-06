import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, fill_color, border_color=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Background color
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def add_header_footer(doc, CRIMSON):
    # Header
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.text = "Ahadu Bank S.C. | eLearning User Manual"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.runs[0].font.color.rgb = CRIMSON
    hp.runs[0].font.size = Pt(9)
    hp.runs[0].font.bold = True
    
    # Footer with Page Numbers
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.text = "Page "
    fp.runs[0].font.color.rgb = RGBColor(113, 128, 150) # #718096
    fp.runs[0].font.size = Pt(9)
    
    # Add page number field
    run = fp.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def create_callout_box(doc, text, bg_color, text_color, is_bold=False):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Set width
    for row in table.rows:
        row.cells[0].width = Inches(6.0)
    
    cell = table.rows[0].cells[0]
    set_cell_background(cell, bg_color)
    p = cell.paragraphs[0]
    p.text = text
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.runs[0].font.color.rgb = text_color
    if is_bold:
        p.runs[0].font.bold = True
    # Add a little spacing after the table
    doc.add_paragraph()

def create_manual(output_path):
    doc = Document()
    
    # Ahadu Bank Brand Colors
    CRIMSON = RGBColor(139, 16, 48)  # #8B1030
    NAVY = RGBColor(12, 35, 64)      # #0C2340
    GOLD = RGBColor(200, 169, 81)    # #C8A951
    WINE = RGBColor(92, 10, 30)      # #5C0A1E
    LIGHT_CRIMSON = "FDF2F4"         
    LIGHT_GOLD = "FFFBF0"
    
    # Base font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(45, 55, 72)
    
    add_header_footer(doc, CRIMSON)
    
    # ================= COVER PAGE =================
    for _ in range(5): doc.add_paragraph()
    
    title = doc.add_heading('AHADU eLEARNING', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = CRIMSON
    title.runs[0].font.size = Pt(42)
    title.runs[0].font.bold = True
    
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('End-to-End User Manual')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.color.rgb = NAVY
    subtitle.runs[0].font.size = Pt(22)
    
    for _ in range(3): doc.add_paragraph()
    
    cover_table = doc.add_table(rows=1, cols=1)
    cover_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = cover_table.rows[0].cells[0]
    set_cell_background(cell, "FDF2F4") # Light crimson bg for the meta box
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.add_run("Module Version: 1.0\n").bold = True
    cp.add_run("Platform: Odoo 18\n").bold = True
    cp.add_run("Prepared for: Ahadu Bank S.C.\n").bold = True
    cp.add_run("Date: March 2026\n").bold = True
    cp.add_run("Classification: Internal Use Only").bold = True
    for r in cp.runs:
        r.font.color.rgb = CRIMSON
        r.font.size = Pt(12)
        
    doc.add_page_break()
    
    # ================= TABLE OF CONTENTS =================
    toc_heading = doc.add_heading('Table of Contents', level=1)
    toc_heading.runs[0].font.color.rgb = CRIMSON
    toc_heading.runs[0].font.size = Pt(24)
    doc.add_paragraph()
    
    chapters = [
        "1. Introduction",
        "2. Getting Started",
        "3. Managing Courses",
        "4. Managing Course Content (Slides)",
        "5. Tags & Course Organization",
        "6. Quizzes & Assessments",
        "7. Enrollment & Learner Progress",
        "8. Best Practices & Tips",
        "9. Glossary"
    ]
    for ch in chapters:
        p = doc.add_paragraph(ch)
        p.runs[0].font.color.rgb = NAVY
        p.runs[0].font.size = Pt(13)
        p.runs[0].font.bold = True
        
    doc.add_page_break()
    
    # Helper to add section headers
    def add_h1(text):
        h = doc.add_heading(text, level=1)
        h.runs[0].font.color.rgb = NAVY
        h.runs[0].font.size = Pt(20)
        h.runs[0].font.bold = True
        
    def add_h2(text):
        h = doc.add_heading(text, level=2)
        h.runs[0].font.color.rgb = CRIMSON
        h.runs[0].font.size = Pt(16)
        
    def add_h3(text):
        h = doc.add_heading(text, level=3)
        h.runs[0].font.color.rgb = WINE
        h.runs[0].font.size = Pt(13)

    # ================= CHAPTER 1 =================
    add_h1('1. Introduction')
    
    add_h2('1.1 About Ahadu eLearning')
    doc.add_paragraph("The Ahadu eLearning module is a customized extension of Odoo 18's eLearning (Website Slides) platform, specifically tailored for Ahadu Bank S.C.. It provides a comprehensive Learning Management System (LMS) that enables the bank to create, manage, and deliver training courses and educational content to employees across all branches.")
    
    add_h2('1.2 Key Features')
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'Description'
    for c in hdr_cells:
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(c, "8B1030")
        
    features = [
        ("Course Management", "Create and organize training courses with sections, categories, and tags"),
        ("Multi-format Content", "Support for videos, documents, infographics, web pages, and articles"),
        ("Local Video Upload", "Upload video files directly (up to 1.2 GB) from your device"),
        ("Google Drive Integration", "Embed videos from Google Drive with automatic player rendering"),
        ("External Link Support", "Link to external training resources and video platforms"),
        ("Quiz & Assessments", "Create quiz questions with multiple answers for knowledge evaluation"),
        ("Course Tagging", "Organize courses using tags and tag groups for easy filtering"),
        ("Progress Tracking", "Monitor employee enrollment and completion status"),
        ("Ahadu Branding", "Complete Ahadu Bank branding throughout the interface")
    ]
    for i, (feat, desc) in enumerate(features):
        row_cells = table.add_row().cells
        row_cells[0].text = feat
        row_cells[1].text = desc
        if i % 2 == 1:
            set_cell_background(row_cells[0], LIGHT_CRIMSON)
            set_cell_background(row_cells[1], LIGHT_CRIMSON)

    doc.add_paragraph('\n')
    add_h2('1.3 System Requirements')
    reqs = ["Odoo 18 Community or Enterprise Edition", "Website Slides (website_slides) module installed", "Modern web browser (Chrome, Firefox, Edge, or Safari)", "Stable internet connection"]
    for r in reqs:
        doc.add_paragraph(r, style='List Bullet')
        
    add_h2('1.4 User Roles')
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdr_cells = table2.rows[0].cells
    for i, text in enumerate(['Role', 'Access Level', 'Description']):
        hdr_cells[i].text = text
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[i], "8B1030")

    roles = [
        ("eLearning Admin", "Full Access", "Create, edit, delete courses and all content. Manage configurations."),
        ("eLearning Manager", "Manager Access", "Create and manage courses, view reports and analytics."),
        ("eLearning Officer", "User Access", "Create content, manage assigned courses."),
        ("Employee (Learner)", "Portal Access", "Enroll in courses, view content, take quizzes.")
    ]
    for i, (r, a, d) in enumerate(roles):
        row_cells = table2.add_row().cells
        row_cells[0].text = r
        row_cells[1].text = a
        row_cells[2].text = d
        if i % 2 == 1:
            for c in row_cells:
                set_cell_background(c, LIGHT_CRIMSON)

    doc.add_page_break()

    # ================= CHAPTER 2 =================
    add_h1('2. Getting Started')
    add_h2('2.1 Accessing Ahadu eLearning')
    doc.add_paragraph("To access the Ahadu eLearning module, follow these steps:")
    create_callout_box(doc, "Step 1: Open your web browser and navigate to the Ahadu ERP URL (e.g., https://erp.ahadubank.com.et).", LIGHT_CRIMSON, CRIMSON)
    create_callout_box(doc, "Step 2: Enter your login credentials (email and password) provided by your system administrator.", LIGHT_CRIMSON, CRIMSON)
    create_callout_box(doc, 'Step 3: Click the "Log in" button to access the main dashboard.', LIGHT_CRIMSON, CRIMSON)
    
    create_callout_box(doc, "📸 [Screenshot Placeholder: Odoo Login Page]\n(Point arrow to Email and Password fields)", LIGHT_CRIMSON, CRIMSON)
    
    create_callout_box(doc, 'Step 4: From the main application menu, locate and click on "Ahadu eLearning" to open the module.', LIGHT_CRIMSON, CRIMSON)
    
    create_callout_box(doc, "📸 [Screenshot Placeholder: Main App Menu — Ahadu eLearning Icon]\n(Point arrow to Ahadu eLearning app)", LIGHT_CRIMSON, CRIMSON)
    
    create_callout_box(doc, 'ℹ️ Note: If you cannot see the Ahadu eLearning app on your dashboard, contact your system administrator to ensure you have the proper access rights assigned to your user account.', LIGHT_GOLD, GOLD, True)
    
    add_h2('2.2 Module Dashboard Overview')
    doc.add_paragraph('Upon opening the Ahadu eLearning module, you will see the main dashboard displaying all available courses in a Kanban view (card layout). Each course card shows:')
    
    dash_items = ['Course Name — The title of the training course', 'Course Image — Visual thumbnail for easy identification', 'Content Count — Number of slides/lessons in the course', 'Enrollment Status — How many participants have enrolled', 'Publication Status — Whether the course is published or in draft']
    for i in dash_items:
        doc.add_paragraph(i, style='List Bullet')

    create_callout_box(doc, "📸 [Screenshot Placeholder: Main Dashboard (Kanban View)]\n(Point arrows to course cards)", LIGHT_CRIMSON, CRIMSON)
    
    add_h2('2.3 Navigation Menu')
    p = doc.add_paragraph('The top navigation bar within the Ahadu eLearning module provides access to:')
    
    table3 = doc.add_table(rows=1, cols=2)
    table3.style = 'Table Grid'
    hdr_cells = table3.rows[0].cells
    for i, text in enumerate(['Menu Item', 'Description']):
        hdr_cells[i].text = text
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[i], "8B1030")

    menus = [
        ("Ahadu Courses -> All Courses", "View and manage all training courses"),
        ("Ahadu Courses -> Contents", "View all slide content across courses"),
        ("Reporting", "View analytics and reports on course performance"),
        ("Configuration -> Tags", "Manage course tags and tag groups"),
        ("Configuration -> Settings", "Configure module settings and preferences")
    ]
    for i, (m, d) in enumerate(menus):
        row_cells = table3.add_row().cells
        row_cells[0].text = m
        row_cells[1].text = d
        if i % 2 == 1:
            for c in row_cells:
                set_cell_background(c, LIGHT_CRIMSON)

    doc.add_paragraph('\n')
    create_callout_box(doc, "📸 [Screenshot Placeholder: Top Navigation Menu Bar]\n(Point arrows to menu items)", LIGHT_CRIMSON, CRIMSON)
    
    doc.add_page_break()
    
    # ================= CHAPTER 3 =================
    add_h1('3. Managing Courses')
    add_h2('3.1 Creating a New Course')
    create_callout_box(doc, "Step 1: Navigate to Ahadu Courses -> All Courses from the top menu.", LIGHT_CRIMSON, CRIMSON)
    create_callout_box(doc, 'Step 2: Click the "New" button located at the top-left corner of the page.', LIGHT_CRIMSON, CRIMSON)
    create_callout_box(doc, "📸 [Screenshot Placeholder: Courses List — New Button]\n(Point arrow to 'New')", LIGHT_CRIMSON, CRIMSON)
    create_callout_box(doc, 'Step 3: Fill in the course details in the form view:', LIGHT_CRIMSON, CRIMSON)
    
    table4 = doc.add_table(rows=1, cols=2)
    table4.style = 'Table Grid'
    hdr_cells = table4.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Description'
    for c in hdr_cells:
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(c, "8B1030")
        
    fields = [
        ("Course Title (Required)", "Enter a descriptive name for the course"),
        ("Website", "Select which website will display the course"),
        ("Responsible", "Assign a user responsible for managing the course"),
        ("Tags", "Add tags to categorize the course (e.g., Compliance, IT)"),
        ("Description", "Provide a detailed description of the course content"),
        ("Course Image", "Upload an image that represents the course")
    ]
    for i, (f, d) in enumerate(fields):
        row_cells = table4.add_row().cells
        row_cells[0].text = f
        row_cells[1].text = d
        if i % 2 == 1:
             set_cell_background(row_cells[0], LIGHT_CRIMSON)
             set_cell_background(row_cells[1], LIGHT_CRIMSON)
             
    doc.add_paragraph('\n')
    create_callout_box(doc, "📸 [Screenshot Placeholder: New Course Form]\n(Point arrows to essential fields)", LIGHT_CRIMSON, CRIMSON)
    create_callout_box(doc, "Step 4: Configure the Options tab to set Enroll Policy, Access Rights, Reviews, etc.", LIGHT_CRIMSON, CRIMSON)
    create_callout_box(doc, "Step 5: Click 'Save' to finalize course creation.", LIGHT_CRIMSON, CRIMSON)
    
    create_callout_box(doc, '💡 Tip: Use meaningful course titles that clearly describe the training topic. For example, use "Anti-Money Laundering (AML) Compliance Training Q1 2026" instead of just "AML Training".', LIGHT_GOLD, GOLD, True)
    
    add_h2('3.2 Adding Course Sections')
    create_callout_box(doc, "Step 1: Open course, navigate to 'Content' tab.", LIGHT_CRIMSON, CRIMSON)
    create_callout_box(doc, "Step 2: Click 'Add a Section' and enter the section name.", LIGHT_CRIMSON, CRIMSON)
    create_callout_box(doc, "📸 [Screenshot Placeholder: Adding a Section]", LIGHT_CRIMSON, CRIMSON)
    
    add_h2('3.3 Publishing a Course')
    create_callout_box(doc, "Step 1: Open the course, click 'Go to Website'.\nStep 2: Toggle 'Published' switch.", LIGHT_CRIMSON, CRIMSON)
    create_callout_box(doc, "📸 [Screenshot Placeholder: Publish Toggle]", LIGHT_CRIMSON, CRIMSON)

    add_h2('3.4 Editing and Deleting Courses')
    doc.add_paragraph("To Edit: Click on any course from the list to open it. Modify the desired fields and save.", style='List Bullet')
    doc.add_paragraph('To Delete: Open the course, click the ⚙️ Actions (gear icon) menu, and select "Delete". Confirm.', style='List Bullet')

    doc.add_page_break()
    
    # ================= CHAPTER 4 =================
    add_h1('4. Managing Course Content (Slides)')
    add_h2('4.1 Content Types Overview')
    
    table5 = doc.add_table(rows=1, cols=2)
    table5.style = 'Table Grid'
    hdr_cells = table5.rows[0].cells
    hdr_cells[0].text = 'Content Type'
    hdr_cells[1].text = 'Use Case'
    for c in hdr_cells:
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(c, "8B1030")
        
    types = [
        ("📄 Document", "PDF documents, policy manuals, procedures"),
        ("🎬 Video", "Training videos, recorded lectures, demonstrations"),
        ("🖼️ Infographic", "Images, charts, visual presentations"),
        ("📝 Article", "Rich text articles written directly in Odoo"),
        ("🌐 Webpage", "Custom HTML-based content pages")
    ]
    for i, (t, u) in enumerate(types):
        row_cells = table5.add_row().cells
        row_cells[0].text = t
        row_cells[1].text = u
        if i % 2 == 1:
            set_cell_background(row_cells[0], LIGHT_CRIMSON)
            set_cell_background(row_cells[1], LIGHT_CRIMSON)
            
    doc.add_paragraph('\n')
    add_h2('4.2 Adding New Content')
    create_callout_box(doc, "Step 1: In the 'Content' tab, click 'Add Content'.", LIGHT_CRIMSON, CRIMSON)
    create_callout_box(doc, "📸 [Screenshot Placeholder: Add Content Button]", LIGHT_CRIMSON, CRIMSON)
    create_callout_box(doc, "Step 2: Fill in Title, Content Type, and Source Type.", LIGHT_CRIMSON, CRIMSON)
    
    add_h2('4.3 Video Source Types (Custom Feature)')
    add_h3('Option A: Upload from Device')
    create_callout_box(doc, "Select 'Video' > 'Upload from device'. Supports files up to 1.2 GB.", LIGHT_CRIMSON, CRIMSON)
    add_h3('Option B: External URL (Google Drive)')
    create_callout_box(doc, "Paste Google Drive URL. System automatically renders the player.", LIGHT_CRIMSON, CRIMSON)
    add_h3('Option C: External Link')
    create_callout_box(doc, "Link to external video platform.", LIGHT_CRIMSON, CRIMSON)
    
    create_callout_box(doc, "📸 [Screenshot Placeholder: Video Source Options]", LIGHT_CRIMSON, CRIMSON)
    
    doc.add_page_break()

    # ================= CHAPTER 5 =================
    add_h1('5. Tags & Course Organization')
    add_h2('5.1 Understanding Tags')
    doc.add_paragraph("Tags in the Ahadu eLearning module help categorize and organize courses for easy discovery.", style='List Bullet')
    doc.add_paragraph("Course Tags: Applied to courses (e.g., Compliance, IT Skills)", style='List Bullet')
    doc.add_paragraph("Content Tags: Applied to individual slides/content within courses", style='List Bullet')
    
    add_h2('5.2 Managing Course Tags')
    create_callout_box(doc, "Step 1: Navigate to Configuration -> Tags.\nStep 2: Click 'New' to create a new tag and associate with a Tag Group.", LIGHT_CRIMSON, CRIMSON)
    create_callout_box(doc, "📸 [Screenshot Placeholder: Tags Menu and Configuration]", LIGHT_CRIMSON, CRIMSON)
    
    doc.add_page_break()

    # ================= CHAPTER 6 =================
    add_h1('6. Quizzes & Assessments')
    add_h2('6.1 Adding Quiz Questions')
    doc.add_paragraph("Each content slide can have quiz questions attached. Learners must answer these correctly to mark the content as completed.")
    create_callout_box(doc, "Step 1: Open a slide and navigate to the 'Quiz' tab.\nStep 2: Click 'Add a line' to create a question.", LIGHT_CRIMSON, CRIMSON)
    
    add_h2('6.2 Configuring Questions')
    create_callout_box(doc, "Enter Question Text, then add Answer lines. Check 'Is Correct' for the right answer.", LIGHT_CRIMSON, CRIMSON)
    create_callout_box(doc, "📸 [Screenshot Placeholder: Quiz Forms & Right/Wrong Answers]", LIGHT_CRIMSON, CRIMSON)

    doc.add_page_break()

    # ================= CHAPTER 7 =================
    add_h1('7. Enrollment & Learner Progress')
    add_h2('7.1 Course Enrollment Policies')
    doc.add_paragraph("Open: Any employee with access can self-enroll", style='List Bullet')
    doc.add_paragraph("On Invitation: Only invited employees can access", style='List Bullet')
    
    add_h2('7.2 Inviting Employees')
    create_callout_box(doc, "Click 'Share' or 'Invite' on the course, enter email addresses, and send the invitation message.", LIGHT_CRIMSON, CRIMSON)
    create_callout_box(doc, "📸 [Screenshot Placeholder: Email Invitation Dialog]", LIGHT_CRIMSON, CRIMSON)
    
    add_h2('7.3 Tracking Progress & Reports')
    create_callout_box(doc, "Go to the 'Attendees' tab to view progress of enrolled employees, or use the 'Reporting' menu for overall analytics.", LIGHT_CRIMSON, CRIMSON)
    create_callout_box(doc, "📸 [Screenshot Placeholder: Reporting Dashboard]", LIGHT_CRIMSON, CRIMSON)

    doc.add_page_break()

    # ================= CHAPTER 8 & 9 =================
    add_h1('8. Best Practices & Tips')
    add_h2('8.1 Course Design')
    doc.add_paragraph("Keep content modular (5-15 min sections)", style='List Bullet')
    doc.add_paragraph("Mix content types (Video, Text, Quiz)", style='List Bullet')
    doc.add_paragraph("Include frequent short quizzes to reinforce knowledge", style='List Bullet')

    add_h2('8.2 Video Upload Guidelines')
    create_callout_box(doc, "Format: MP4 (H.264)\nMax Size: 1.2 GB\nResolution: 720p Recommended", LIGHT_CRIMSON, CRIMSON)

    add_h1('9. Glossary')
    table6 = doc.add_table(rows=1, cols=2)
    table6.style = 'Table Grid'
    hdr_cells = table6.rows[0].cells
    hdr_cells[0].text = 'Term'
    hdr_cells[1].text = 'Definition'
    for c in hdr_cells:
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(c, "8B1030")

    glossary = [
        ("Course (Channel)", "A collection of training content"),
        ("Slide / Content", "An individual piece of learning material"),
        ("Section", "A grouping header to organize slides"),
        ("Kanban View", "Card-based visual layout for browsing courses")
    ]
    for i, (t, d) in enumerate(glossary):
        row_cells = table6.add_row().cells
        row_cells[0].text = t
        row_cells[1].text = d
        if i % 2 == 1:
            set_cell_background(row_cells[0], LIGHT_CRIMSON)
            set_cell_background(row_cells[1], LIGHT_CRIMSON)

    doc.save(output_path)

if __name__ == '__main__':
    create_manual('/opt/odoo18/odoo18/custom_addons/ahadu_elearning/Ahadu_eLearning_User_Manual.docx')
