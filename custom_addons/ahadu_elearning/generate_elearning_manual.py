#!/usr/bin/env python3
"""
Ahadu Bank eLearning Module — Full User Manual Generator
Produces a comprehensive, branded DOCX manual.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────
# BRAND COLORS  (Light Burgundy / Ahadu Bank palette)
# ─────────────────────────────────────────────────────────────
BURGUNDY      = RGBColor(128,  0,  32)   # #800020  primary brand
DARK_BURGUNDY = RGBColor( 80,  0,  20)   # #500014  deep accent
NAVY          = RGBColor( 12, 35,  64)   # #0C2340  secondary
GOLD          = RGBColor(200, 169, 81)   # #C8A951  accent
LIGHT_BG      = "FBF0F3"               # very light pink/burgundy bg
LIGHT_GOLD_BG = "FFFBF0"              # light gold bg for tips
HEADER_HEX    = "800020"              # table header cell fill


# ─────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────
def set_cell_bg(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "Ahadu Bank S.C.  |  eLearning Module User Manual"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.runs[0]
    run.font.color.rgb = BURGUNDY
    run.font.size = Pt(9)
    run.font.bold = True

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.clear()

    r1 = fp.add_run("Confidential  •  Internal Use Only  •  Page ")
    r1.font.color.rgb = RGBColor(113, 128, 150)
    r1.font.size = Pt(8)

    pn = fp.add_run()
    for tag, attrs in [
        ('w:fldChar',   {'w:fldCharType': 'begin'}),
        ('w:instrText', None),
        ('w:fldChar',   {'w:fldCharType': 'separate'}),
        ('w:fldChar',   {'w:fldCharType': 'end'}),
    ]:
        el = OxmlElement(tag)
        if attrs:
            for k, v in attrs.items():
                el.set(qn(k), v)
        if tag == 'w:instrText':
            el.set(qn('xml:space'), 'preserve')
            el.text = 'PAGE'
        pn._r.append(el)

    r2 = fp.add_run("  |  © 2026 Ahadu Bank S.C.")
    r2.font.color.rgb = RGBColor(113, 128, 150)
    r2.font.size = Pt(8)


def styled_table(doc, headers, rows, even_bg=LIGHT_BG):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        c.text = h
        set_cell_bg(c, HEADER_HEX)
        r = c.paragraphs[0].runs[0]
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = tbl.rows[ri + 1].cells[ci]
            c.text = val
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 1:
                set_cell_bg(c, even_bg)
    doc.add_paragraph()
    return tbl


def callout(doc, text, bg=LIGHT_BG, color=None, bold=False):
    color = color or BURGUNDY
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.color.rgb = color
    run.font.size = Pt(10)
    run.font.bold = bold
    doc.add_paragraph()


def tip(doc, text):
    callout(doc, f"💡 Tip:  {text}", bg=LIGHT_GOLD_BG, color=GOLD, bold=True)


def note(doc, text):
    callout(doc, f"ℹ️  Note:  {text}", bg=LIGHT_BG, color=BURGUNDY)


def step(doc, num, text):
    callout(doc, f"Step {num}:  {text}", bg=LIGHT_BG, color=BURGUNDY)


def add_h1(doc, text):
    h = doc.add_heading(text, level=1)
    h.runs[0].font.color.rgb = NAVY
    h.runs[0].font.size = Pt(20)
    h.runs[0].font.bold = True


def add_h2(doc, text):
    h = doc.add_heading(text, level=2)
    h.runs[0].font.color.rgb = BURGUNDY
    h.runs[0].font.size = Pt(15)


def add_h3(doc, text):
    h = doc.add_heading(text, level=3)
    h.runs[0].font.color.rgb = DARK_BURGUNDY
    h.runs[0].font.size = Pt(12)


def body(doc, text):
    p = doc.add_paragraph(text)
    if p.runs:
        p.runs[0].font.size = Pt(11)
    return p


def bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')


# ─────────────────────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────────────────────
def add_cover(doc):
    for _ in range(6):
        doc.add_paragraph()

    t = doc.add_heading('AHADU eLEARNING', 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.color.rgb = BURGUNDY
    t.runs[0].font.size = Pt(40)
    t.runs[0].font.bold = True

    s = doc.add_paragraph('End-to-End User Manual')
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.runs[0].font.color.rgb = NAVY
    s.runs[0].font.size = Pt(22)

    doc.add_paragraph()
    div = doc.add_paragraph('─' * 55)
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div.runs[0].font.color.rgb = GOLD
    div.runs[0].font.size = Pt(14)
    for _ in range(3):
        doc.add_paragraph()

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, LIGHT_BG)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        "Module Version: 1.0.0",
        "Platform: Odoo 18",
        "Prepared for: Ahadu Bank S.C.",
        "Date: March 2026",
        "Classification: Confidential — Internal Use Only",
    ]:
        r = p.add_run(line + "\n")
        r.bold = True
        r.font.color.rgb = BURGUNDY
        r.font.size = Pt(12)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
# TABLE OF CONTENTS
# ─────────────────────────────────────────────────────────────
def add_toc(doc):
    add_h1(doc, 'Table of Contents')
    doc.add_paragraph()
    chapters = [
        "1.  Introduction",
        "2.  Getting Started",
        "3.  Module Dashboard Overview",
        "4.  Managing Courses",
        "5.  Managing Course Content (Slides)",
        "6.  Video Content — Upload Options",
        "7.  Tags & Course Organization",
        "8.  Quizzes & Assessments",
        "9.  Enrollment & Learner Progress",
        "10. Reporting & Analytics",
        "11. Configuration & Settings",
        "12. Publishing Courses",
        "13. Roles & Access Control",
        "14. Best Practices & Tips",
        "15. Glossary",
    ]
    for ch in chapters:
        p = doc.add_paragraph(ch)
        if p.runs:
            p.runs[0].font.color.rgb = NAVY
            p.runs[0].font.size = Pt(12)
            p.runs[0].font.bold = True
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
# CHAPTER 1 — INTRODUCTION
# ─────────────────────────────────────────────────────────────
def ch1_introduction(doc):
    add_h1(doc, '1. Introduction')

    add_h2(doc, '1.1 About Ahadu eLearning')
    body(doc, (
        "The Ahadu eLearning module is a customized Learning Management System (LMS) built on Odoo 18's "
        "Website Slides (website_slides) platform, tailored exclusively for Ahadu Bank S.C. "
        "It enables the bank to create, manage, and deliver training courses and educational content "
        "to employees across all branches — from the head office to the most remote locations."
    ))
    body(doc, (
        "The module extends the standard Odoo LMS with powerful custom capabilities: local video upload "
        "(up to 1.2 GB), Google Drive video embedding with automatic player rendering, external link support, "
        "and complete Ahadu Bank branding throughout the interface."
    ))

    add_h2(doc, '1.2 Key Features')
    styled_table(doc, ['Feature', 'Description'], [
        ("Course Management",       "Create and organize training courses with channels, sections, categories, and tags"),
        ("Multi-format Content",    "Support for videos, PDF documents, infographics, web pages, and rich-text articles"),
        ("Local Video Upload",      "Upload video files directly (MP4, AVI, MKV, etc.) up to 1.2 GB per file"),
        ("Google Drive Integration","Embed Google Drive videos with automatic iframe player rendering"),
        ("External Link Support",   "Link to external training resources and video platforms (YouTube, Vimeo, etc.)"),
        ("Quiz & Assessments",      "Create multiple-choice quiz questions with correct/incorrect answer tracking"),
        ("Course Tagging",          "Organize courses using tags and tag groups for easy discovery and filtering"),
        ("Progress Tracking",       "Monitor employee enrollment, lesson completions, and quiz scores"),
        ("Ahadu Branding",          "Complete Ahadu Bank branding throughout the eLearning interface"),
        ("Reporting",               "Built-in analytics on course performance, enrollments, and completion rates"),
    ])

    add_h2(doc, '1.3 System Requirements')
    for req in [
        "Odoo 18 Community or Enterprise Edition",
        "Website Slides (website_slides) module installed",
        "ahadu_elearning module installed and enabled",
        "Modern web browser: Google Chrome, Firefox, Microsoft Edge, or Safari",
        "Stable internet connection (minimum 2 Mbps for video streaming)",
        "For video uploads: Sufficient server disk space (1.2 GB per video file)",
    ]:
        bullet(doc, req)

    add_h2(doc, '1.4 User Roles & Access')
    styled_table(doc, ['Role', 'Access Level', 'Capabilities'], [
        ("eLearning Administrator", "Full Access",    "Create, edit, delete courses and all content. Manage configurations, users, and settings."),
        ("eLearning Manager",       "Manager Access", "Create and manage courses, slides, and tags. View all reports and analytics."),
        ("eLearning Officer",       "User Access",    "Create content, manage assigned courses, upload videos, add quiz questions."),
        ("Employee (Learner)",      "Portal Access",  "Enroll in published courses, view content, take quizzes, track own progress."),
    ])

    note(doc, "If you cannot see the Ahadu eLearning app on your dashboard, contact your system administrator to ensure you have the proper access rights assigned to your user account.")
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
# CHAPTER 2 — GETTING STARTED
# ─────────────────────────────────────────────────────────────
def ch2_getting_started(doc):
    add_h1(doc, '2. Getting Started')

    add_h2(doc, '2.1 Logging In')
    step(doc, 1, "Open your web browser and navigate to the Ahadu ERP URL (e.g., https://erp.ahadubank.com.et).")
    step(doc, 2, "Enter your assigned Email address and Password in the login form.")
    step(doc, 3, "Click 'Log in'. You will be directed to the main Odoo application dashboard.")

    add_h2(doc, '2.2 Opening the eLearning Module')
    step(doc, 1, "From the main application menu (home screen), locate the 'Ahadu eLearning' application icon.")
    step(doc, 2, "Click the icon to open the eLearning module.")
    step(doc, 3, "The system opens the Course (Channel) list as the default landing view.")
    note(doc, "If the Ahadu eLearning app is not visible on your home screen, your administrator has not granted you eLearning access. Contact the IT/ERP support team.")

    add_h2(doc, '2.3 Navigation Menu Structure')
    styled_table(doc, ['Menu', 'Sub-Menu', 'Description'], [
        ("Ahadu eLearning", "—",            "Root application menu"),
        ("Ahadu Courses",   "All Courses",  "View and manage all training courses (Kanban and List views)"),
        ("Ahadu Courses",   "Contents",     "View and manage all individual slides/lessons across all courses"),
        ("Reporting",       "—",            "Access course analytics, enrollment stats, and completion reports"),
        ("Configuration",   "Tags",         "Manage course tags and tag groups"),
        ("Configuration",   "Settings",     "Configure module-level settings and preferences"),
    ])
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
# CHAPTER 3 — DASHBOARD OVERVIEW
# ─────────────────────────────────────────────────────────────
def ch3_dashboard(doc):
    add_h1(doc, '3. Module Dashboard Overview')

    add_h2(doc, '3.1 Default Kanban View')
    body(doc, (
        "When you open Ahadu eLearning, the default view is a Kanban card layout showing all available courses. "
        "Each course is displayed as a card with the following information:"
    ))
    for item in [
        "Course Image — A visual thumbnail for easy identification",
        "Course Name — The title of the training course",
        "Content Count — Number of slides/lessons inside the course",
        "Enrollment Count — How many employees have joined the course",
        "Publication Status — Whether the course is Published (visible to learners) or Unpublished (draft)",
        "Responsible User — The staff member managing this course",
    ]:
        bullet(doc, item)

    add_h2(doc, '3.2 Switching Views')
    body(doc, "You can switch between Kanban (card layout) and List view using the view toggle icons at the top right of the courses page. Use List view when you need to quickly scan or sort multiple courses.")

    add_h2(doc, '3.3 Search & Filters')
    styled_table(doc, ['Search Option', 'Purpose'], [
        ("Search by Name",      "Type the course title to find a specific course"),
        ("Filter: Published",   "Show only courses currently visible to learners"),
        ("Filter: Unpublished", "Show courses in draft/preparation stage"),
        ("Group by: Responsible", "Group courses by the staff member managing them"),
        ("Group by: Tag",       "Group courses by category tags"),
    ])

    tip(doc, "Use the 'Group by: Tag' option when you manage many courses — it helps you quickly locate all courses in a specific category such as Compliance, IT, or Leadership.")
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
# CHAPTER 4 — MANAGING COURSES
# ─────────────────────────────────────────────────────────────
def ch4_courses(doc):
    add_h1(doc, '4. Managing Courses')

    body(doc, (
        "In Ahadu eLearning, a 'Course' (also called a 'Channel' in Odoo's technical layer) is the top-level "
        "container for all training content. It groups related lessons, sections, and quizzes into a single "
        "structured learning experience."
    ))

    add_h2(doc, '4.1 Creating a New Course')
    step(doc, 1, "Navigate to Ahadu Courses → All Courses.")
    step(doc, 2, "Click the 'New' button at the top-left of the page.")
    step(doc, 3, "The course form opens. Fill in the required and optional fields (see table below).")
    step(doc, 4, "Navigate to the 'Options' tab to configure enrollment policy and access rights.")
    step(doc, 5, "Click 'Save manually' (or Odoo auto-saves) to create the course.")

    add_h2(doc, '4.2 Course Form Fields')
    styled_table(doc, ['Field', 'Required?', 'Description'], [
        ("Course Title",    "✔ Yes",  "A clear, descriptive name for the course (e.g., 'Anti-Money Laundering Compliance Training 2026')"),
        ("Website",         "Optional","The website on which this course will be displayed"),
        ("Responsible",     "Optional","The Ahadu Bank staff member responsible for managing this course"),
        ("Tags",            "Optional","One or more category tags applied to the course for filtering"),
        ("Course Image",    "Optional","An image or logo that represents the course, shown on the Kanban card"),
        ("Description",     "Optional","A full description of what the course covers — shown to learners on the course page"),
        ("Target Audience", "Optional","Short note describing who the course is designed for"),
    ])

    add_h2(doc, '4.3 Options Tab Settings')
    styled_table(doc, ['Setting', 'Description'], [
        ("Enroll Policy",    "Controls who can enroll: 'Public' (anyone), 'On Invitation' (email invite required), or 'On Payment'"),
        ("Access Rights",    "Defines visibility — Website (shows on portal), Internal (employees only)"),
        ("Reviews",          "Allow or disallow learner course reviews and star ratings"),
        ("Certificate",      "Whether to issue a completion certificate upon finishing the course"),
        ("Forum",            "Enable a discussion forum linked to this course for Q&A"),
        ("Karma",            "Points earned by viewing/completing this course (gamification)"),
    ])

    add_h2(doc, '4.4 Adding Sections to a Course')
    body(doc, "Sections are organizational dividers within a course that group related slides/lessons together (e.g., 'Week 1: Introduction', 'Week 2: Core Topics').")
    step(doc, 1, "Open the course, navigate to the 'Content' tab.")
    step(doc, 2, "Click 'Add a Section' at the bottom of the content list.")
    step(doc, 3, "Type the section name (e.g., 'Module 1 — Fundamentals') and press Enter.")
    step(doc, 4, "Drag and drop lessons under each section to organize your course structure.")

    add_h2(doc, '4.5 Editing a Course')
    step(doc, 1, "Open the course from Ahadu Courses → All Courses.")
    step(doc, 2, "Click into any field and make your changes.")
    step(doc, 3, "Odoo auto-saves changes. Click 'Save manually' if needed.")

    add_h2(doc, '4.6 Deleting a Course')
    step(doc, 1, "Open the course you wish to delete.")
    step(doc, 2, "Click the ⚙️ Actions (gear) icon in the top toolbar.")
    step(doc, 3, "Select 'Delete' from the dropdown menu.")
    step(doc, 4, "A confirmation dialog appears — click 'OK' to permanently delete the course.")
    callout(doc, "⚠️ Warning: Deleting a course permanently removes all its content (slides, quizzes, enrollments). This action cannot be undone. Consider archiving the course instead.", bg=LIGHT_GOLD_BG, color=GOLD, bold=True)

    add_h2(doc, '4.7 Archiving a Course')
    step(doc, 1, "Open the course, click the ⚙️ Actions menu.")
    step(doc, 2, "Select 'Archive'. The course will be hidden from all views but can be restored later.")
    tip(doc, "Always archive rather than delete courses that may need to be reused or audited in the future.")
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
# CHAPTER 5 — MANAGING CONTENT (SLIDES)
# ─────────────────────────────────────────────────────────────
def ch5_content(doc):
    add_h1(doc, '5. Managing Course Content (Slides)')

    body(doc, (
        "A 'Slide' is the individual unit of learning content within a course. Each slide represents one lesson, "
        "video, document, infographic, article, or webpage. Slides are organized within a course under sections."
    ))

    add_h2(doc, '5.1 Content Types Overview')
    styled_table(doc, ['Content Type', 'Icon', 'Best Used For'], [
        ("Video",      "🎬", "Training videos, recorded lectures, demonstrations, screen recordings"),
        ("Document",   "📄", "PDF policy manuals, procedures, guidelines, forms"),
        ("Infographic","🖼️", "Images, charts, visual summaries, diagrams"),
        ("Article",    "📝", "Rich-text articles written directly inside Odoo using the HTML editor"),
        ("Web Page",   "🌐", "Custom HTML-based content pages with embedded elements"),
    ])

    add_h2(doc, '5.2 Adding New Content to a Course')
    step(doc, 1, "Open the course and navigate to the 'Content' tab.")
    step(doc, 2, "Click 'Add Content' below the course section where you want to add the lesson.")
    step(doc, 3, "A new slide form (popup or inline) appears. Fill in the Title.")
    step(doc, 4, "Select the 'Slide Category' (Video, Document, Infographic, Article, or Web Page).")
    step(doc, 5, "Set the content source based on the selected content type (see Chapter 6 for video options).")
    step(doc, 6, "Configure optional fields: Tags, Description, Quiz questions.")
    step(doc, 7, "Click 'Save' or 'Add' to finalize the slide.")

    add_h2(doc, '5.3 Slide Form Fields')
    styled_table(doc, ['Field', 'Description'], [
        ("Title",           "The name of the slide/lesson shown to learners"),
        ("Course",          "The course this slide belongs to (auto-filled when adding from course)"),
        ("Slide Category",  "Content type: Video, Document, Infographic, Article, or Web Page"),
        ("Sequence",        "The display order within the course (lower number = shown first)"),
        ("Tags",            "Content-level tags for the slide (different from course tags)"),
        ("Completion",      "Whether this slide requires quiz completion to mark as done"),
        ("Completion Time", "Estimated time (in minutes) for a learner to complete this slide"),
        ("Description",     "Brief description of what this slide covers"),
        ("Quiz Tab",        "Add quiz questions learners must answer to complete this slide"),
    ])

    add_h2(doc, '5.4 Editing and Reordering Slides')
    body(doc, "To edit a slide, click on it from the course Content tab. To reorder, drag and drop slides using the handle icon (⠿) on the left side of each slide row.")

    add_h2(doc, '5.5 Deleting a Slide')
    step(doc, 1, "Open the course and navigate to the Content tab.")
    step(doc, 2, "Click the slide you want to remove.")
    step(doc, 3, "Click the ⚙️ Actions menu → 'Delete' and confirm.")
    note(doc, "Deleting a slide also removes all quiz questions and learner completion records tied to that slide.")
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
# CHAPTER 6 — VIDEO CONTENT OPTIONS
# ─────────────────────────────────────────────────────────────
def ch6_video(doc):
    add_h1(doc, '6. Video Content — Upload Options')

    body(doc, (
        "The Ahadu eLearning module offers three distinct methods for incorporating video content into courses. "
        "Each method is suited to a different scenario depending on file size, storage location, and access requirements."
    ))

    add_h2(doc, '6.1 Option A — Upload from Device (Local File)')
    body(doc, "This option allows you to upload a video file directly from your computer to the Odoo server. The video is stored within the Ahadu Bank ERP system itself.")
    step(doc, 1, "In the slide form, set Slide Category to 'Video'.")
    step(doc, 2, "Set Source Type to 'Upload from device'.")
    step(doc, 3, "Click the 'Upload Video' field to browse and select your video file.")
    step(doc, 4, "Wait for the upload to complete (progress shown in browser).")
    step(doc, 5, "The system generates a secure playback URL and auto-detects the MIME type.")
    step(doc, 6, "Save the slide. The video is immediately playable by enrolled learners.")

    styled_table(doc, ['Specification', 'Value'], [
        ("Maximum File Size",  "1.2 GB per video"),
        ("Supported Formats",  "MP4 (recommended), AVI, MKV, MOV, WebM"),
        ("Recommended Format", "MP4 with H.264 video codec and AAC audio codec"),
        ("Resolution",         "720p (1280×720) recommended; 1080p supported"),
        ("Player",             "Built-in HTML5 <video> player; no download allowed"),
        ("Access Control",     "Auto-set to Public so enrolled learners can stream"),
    ])

    tip(doc, "For best streaming performance on the Ahadu network, compress videos to 720p and target 500 kbps to 2 Mbps bitrate before uploading.")

    add_h2(doc, '6.2 Option B — Google Drive Embed (External / URL)')
    body(doc, "This option allows you to link a video already stored in Google Drive. The system automatically extracts the file ID and renders the video using Google Drive's preview iframe player.")
    step(doc, 1, "Upload/share your video on Google Drive and set it to 'Anyone with the link can view'.")
    step(doc, 2, "Copy the Google Drive sharing link (e.g., https://drive.google.com/file/d/FILE_ID/view).")
    step(doc, 3, "In the slide form, set Slide Category to 'Video' and Source Type to 'External (URL)'.")
    step(doc, 4, "Paste the Google Drive link into the 'Video URL' field.")
    step(doc, 5, "The system auto-detects the file ID and converts it to an embed URL. No manual extraction needed.")
    step(doc, 6, "Save the slide. Learners will see the Google Drive video player.")

    note(doc, "If the Google Drive file is restricted or 'Ahadu Bank only', learners who are not logged into Google with an authorized account will see an access error. Use 'Anyone with the link' sharing for full compatibility.")

    add_h2(doc, '6.3 Option C — External Link')
    body(doc, "This option allows you to link to any external video or webpage URL (e.g., YouTube, Vimeo, a public training website). When a learner clicks the slide, they are directed to the external URL in a new browser tab.")
    step(doc, 1, "Set Slide Category to 'Video' and Source Type to 'External Link'.")
    step(doc, 2, "Enter the full external URL in the Video URL field.")
    step(doc, 3, "Save the slide.")
    tip(doc, "Use the External Link option sparingly — it directs learners away from the Ahadu eLearning platform, making it harder to track completion and quiz results.")

    add_h2(doc, '6.4 Source Type Comparison Summary')
    styled_table(doc, ['Feature', 'Upload from Device', 'Google Drive', 'External Link'], [
        ("Storage Location",    "Ahadu ERP Server",    "Google Drive",         "External website"),
        ("Max File Size",       "1.2 GB",              "No limit (Drive)",     "N/A"),
        ("Requires Internet",   "Yes (for streaming)", "Yes",                  "Yes"),
        ("Access Control",      "Managed in Odoo",     "Managed in Drive",     "None (public URL)"),
        ("Auto Player Render",  "Yes (HTML5)",         "Yes (Drive preview)",  "No (redirects away)"),
        ("Download by Learner", "No (disabled)",       "Depends on Drive rights","Yes"),
        ("Recommended for",     "Official training videos stored internally", "Large existing Drive library", "External partner content"),
    ])
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
# CHAPTER 7 — TAGS & ORGANIZATION
# ─────────────────────────────────────────────────────────────
def ch7_tags(doc):
    add_h1(doc, '7. Tags & Course Organization')

    body(doc, (
        "Tags are labels applied to courses and slides to make them easier to categorize, search, and filter. "
        "They are especially useful when you have a large number of courses and need to quickly find content by topic."
    ))

    add_h2(doc, '7.1 Types of Tags')
    styled_table(doc, ['Tag Type', 'Applied To', 'Example'], [
        ("Course Tag",   "Entire course",          "Compliance, IT Skills, Leadership, Teller Operations"),
        ("Tag Group",    "Groups of course tags",  "Department-Level (groups IT, Operations, HR tags together)"),
        ("Content Tag",  "Individual slides",      "Policy Update, Video Lesson, Quick Reference"),
    ])

    add_h2(doc, '7.2 Managing Course Tags')
    step(doc, 1, "Navigate to Configuration → Tags.")
    step(doc, 2, "The tags list shows all existing tags. Click on a tag to edit it.")
    step(doc, 3, "Click 'New' to create a new tag.")
    step(doc, 4, "Enter the Tag Name and assign it a Tag Group (optional but recommended for organization).")
    step(doc, 5, "Save the tag. It is immediately available to assign to courses.")

    add_h2(doc, '7.3 Assigning Tags to Courses')
    step(doc, 1, "Open a course and click in the 'Tags' field on the course form.")
    step(doc, 2, "Start typing the tag name and select from the dropdown.")
    step(doc, 3, "Multiple tags can be assigned to a single course.")
    step(doc, 4, "Save the course.")

    tip(doc, "Create a consistent tagging taxonomy before creating lots of courses. For example: use department names (HR, IT, Finance, Operations) as top-level tag groups, and topic areas as tags within each group.")
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
# CHAPTER 8 — QUIZZES
# ─────────────────────────────────────────────────────────────
def ch8_quizzes(doc):
    add_h1(doc, '8. Quizzes & Assessments')

    body(doc, (
        "Built-in quiz functionality allows course creators to attach multiple-choice questions to any slide. "
        "Learners must answer these questions correctly to mark the slide as completed. "
        "Quizzes are the primary mechanism for knowledge verification within the Ahadu eLearning platform."
    ))

    add_h2(doc, '8.1 Adding Quiz Questions to a Slide')
    step(doc, 1, "Open a course and click on the slide you want to add quiz questions to.")
    step(doc, 2, "Navigate to the 'Quiz' tab in the slide form.")
    step(doc, 3, "Click 'Add a line' to create a new question.")
    step(doc, 4, "In the Question field, type the full question text.")
    step(doc, 5, "Click 'Save & Close' on the question line.")

    add_h2(doc, '8.2 Adding Answer Options')
    step(doc, 1, "Open the question line by clicking on it (or the expand icon ▶).")
    step(doc, 2, "In the Answers section, click 'Add a line' for each answer option.")
    step(doc, 3, "Type each answer text.")
    step(doc, 4, "Check the 'Is Correct' checkbox for the correct answer. Only one answer can be marked correct per question.")
    step(doc, 5, "Save the answers and close the question form.")

    add_h2(doc, '8.3 Quiz Configuration Options')
    styled_table(doc, ['Setting', 'Description'], [
        ("Allow Multiple Attempts", "Whether learner can retry the quiz if they answer incorrectly"),
        ("Quiz Karma Gain",         "Points the learner earns for correctly answering this slide's quiz"),
        ("Completion Mode",         "Whether quiz is required to mark the slide as 'Done'"),
    ])

    add_h2(doc, '8.4 Viewing Quiz Results')
    body(doc, "Quiz results are visible to eLearning Managers and Officers through the Reporting menu. Individual learner scores can be viewed in the Attendees tab of each course.")

    note(doc, "There is no automatic time limit on quizzes in the standard module. If time-limited assessments are needed, contact the ERP development team for a custom solution.")
    tip(doc, "Add at least 3-5 quiz questions per slide for meaningful knowledge checks. Mix fact-recall and application-based questions for best results.")
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
# CHAPTER 9 — ENROLLMENT & PROGRESS
# ─────────────────────────────────────────────────────────────
def ch9_enrollment(doc):
    add_h1(doc, '9. Enrollment & Learner Progress')

    add_h2(doc, '9.1 Enrollment Policies')
    styled_table(doc, ['Policy', 'Description', 'When to Use'], [
        ("Public",         "Any logged-in user can self-enroll by visiting the course page",   "General awareness training open to all employees"),
        ("On Invitation",  "Only employees who receive an email invitation can enroll",        "Targeted mandatory training for specific teams/branches"),
        ("On Payment",     "Enrollment requires a payment (rarely used in bank context)",      "External or special certification programs (if applicable)"),
    ])

    add_h2(doc, '9.2 Setting Enrollment Policy')
    step(doc, 1, "Open the course, navigate to the 'Options' tab.")
    step(doc, 2, "In the 'Enroll Policy' field, select the desired policy.")
    step(doc, 3, "Save the course.")

    add_h2(doc, '9.3 Inviting Employees to a Course')
    step(doc, 1, "Open the course. Click the 'Share' or '📧 Invite' button in the top area of the course form.")
    step(doc, 2, "In the invitation dialog, enter the employee email addresses (or select from Odoo contacts).")
    step(doc, 3, "Customize the invitation message if needed.")
    step(doc, 4, "Click 'Send'. Each recipient receives an email with a direct course enrollment link.")

    add_h2(doc, '9.4 Monitoring Learner Progress')
    step(doc, 1, "Open the course and navigate to the 'Attendees' tab.")
    step(doc, 2, "The tab shows a list of all enrolled learners with their completion percentage.")
    step(doc, 3, "Click on an individual learner to see which slides they have completed and which are pending.")

    add_h2(doc, '9.5 Completion Status Indicators')
    styled_table(doc, ['Status', 'Meaning'], [
        ("In Progress (0-99%)", "Learner has started but not finished all slides"),
        ("Completed (100%)",    "Learner has viewed all required slides and passed all quizzes"),
        ("Not Started (0%)",    "Learner enrolled but has not viewed any content yet"),
    ])

    tip(doc, "Send reminder emails to employees who are 'Not Started' or have been 'In Progress' for more than 2 weeks, especially for mandatory compliance training.")
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
# CHAPTER 10 — REPORTING
# ─────────────────────────────────────────────────────────────
def ch10_reporting(doc):
    add_h1(doc, '10. Reporting & Analytics')

    body(doc, (
        "The Ahadu eLearning reporting section provides analytics on course performance, learner engagement, "
        "and content effectiveness. Reports are accessible from the 'Reporting' menu."
    ))

    add_h2(doc, '10.1 Available Reports')
    styled_table(doc, ['Report', 'What It Shows'], [
        ("Courses Overview",          "All courses with enrollment count, completion rate, and publication status"),
        ("Slide Analysis",            "View and completion counts per slide across all courses"),
        ("Attendees / Enrolments",    "List of learners enrolled in each course with their completion %"),
        ("Quiz Participation",        "Number of quiz attempts and average score per slide"),
        ("Completion by Employee",    "Which employees have completed which courses (useful for compliance audit)"),
    ])

    add_h2(doc, '10.2 Using the Reporting Menu')
    step(doc, 1, "Navigate to Reporting in the top menu bar.")
    step(doc, 2, "Select the desired report view.")
    step(doc, 3, "Use the search and filter bar to narrow down by Course, Department, Date range, or Status.")
    step(doc, 4, "Click the download icon (📥) or use Action → Export to export data to Excel.")

    add_h2(doc, '10.3 Key Metrics to Monitor')
    for metric in [
        "Enrollment Rate — Percentage of target employees who have enrolled in mandatory courses",
        "Completion Rate — Percentage of enrolled employees who have fully completed the course",
        "Quiz Average Score — Average score across all learners on a quiz-enabled slide",
        "Time to Completion — Average days from enrollment to completion",
        "Drop-off Rate — Learners who enrolled but stopped mid-course without completing",
    ]:
        bullet(doc, metric)

    tip(doc, "Run the 'Completion by Employee' report before audit events or compliance deadlines to identify and follow up with employees who have not yet completed mandatory training.")
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
# CHAPTER 11 — CONFIGURATION
# ─────────────────────────────────────────────────────────────
def ch11_configuration(doc):
    add_h1(doc, '11. Configuration & Settings')

    note(doc, "Configuration options are only accessible to users with eLearning Administrator or Manager access.")

    add_h2(doc, '11.1 Tags Configuration')
    body(doc, "Navigate to Configuration → Tags to manage the tag taxonomy used across all courses.")
    styled_table(doc, ['Field', 'Description'], [
        ("Tag Name",    "The label shown on courses and in search filters"),
        ("Tag Group",   "A grouping category that organizes related tags together"),
    ])
    body(doc, "Best practice: create Tag Groups that match major training domains (e.g., 'Compliance', 'Operations', 'IT', 'Leadership') before creating individual tags.")

    add_h2(doc, '11.2 Module Settings')
    step(doc, 1, "Navigate to Configuration → Settings.")
    step(doc, 2, "The settings page shows global eLearning configuration options.")

    add_h2(doc, '11.3 Available Settings')
    styled_table(doc, ['Setting', 'Description'], [
        ("Forum",            "Enable discussion forums linked to courses"),
        ("Certificate",      "Enable course completion certificates"),
        ("Live Sessions",    "Enable scheduling of live Q&A or webinar sessions"),
        ("Reviews & Ratings","Allow learners to rate and review courses"),
        ("Karma",            "Enable gamification points for course completion activities"),
    ])

    tip(doc, "Enable 'Certificate' and 'Reviews' features for high-value compliance courses. Certificates serve as proof of completion for audit purposes.")
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
# CHAPTER 12 — PUBLISHING
# ─────────────────────────────────────────────────────────────
def ch12_publishing(doc):
    add_h1(doc, '12. Publishing Courses')

    body(doc, (
        "By default, newly created courses are 'Unpublished' — they are not visible to learners on the portal. "
        "A course must be explicitly published before employees can enroll and access its content."
    ))

    add_h2(doc, '12.1 Publishing a Course')
    step(doc, 1, "Open the course from Ahadu Courses → All Courses.")
    step(doc, 2, "Click the 'Go to Website' button to preview the course on the learner portal.")
    step(doc, 3, "On the website preview, locate the 'Published / Unpublished' toggle switch at the top.")
    step(doc, 4, "Click the toggle to switch from 'Unpublished' to 'Published'.")
    step(doc, 5, "The course is now live and visible to all users with the appropriate access level.")

    add_h2(doc, '12.2 Unpublishing a Course')
    step(doc, 1, "Go to the course website preview (click 'Go to Website').")
    step(doc, 2, "Click the 'Published' toggle to switch it back to 'Unpublished'.")
    step(doc, 3, "The course is immediately hidden from the learner portal but remains in the backend.")

    add_h2(doc, '12.3 Publishing Individual Slides')
    body(doc, "Individual slides within a course can also be published or unpublished independently. This allows you to release course content incrementally (e.g., new lesson each week).")
    step(doc, 1, "Open a slide from the course Content tab.")
    step(doc, 2, "In the slide form, look for the 'Is Published' toggle at the top.")
    step(doc, 3, "Toggle it to Published or Unpublished as needed.")

    callout(doc, "⚠️  Note: A COURSE must be published for learners to access it, regardless of whether individual slides are published. Always publish the course first.", bg=LIGHT_GOLD_BG, color=GOLD, bold=True)
    tip(doc, "Before publishing, always preview the course as a learner to verify all videos load correctly, all quiz questions are complete, and the content order is logical.")
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
# CHAPTER 13 — ROLES & ACCESS
# ─────────────────────────────────────────────────────────────
def ch13_roles(doc):
    add_h1(doc, '13. Roles & Access Control')

    body(doc, "Access to the Ahadu eLearning module is controlled through Odoo security groups. Your system administrator assigns these roles from the Settings → Users menu.")

    add_h2(doc, '13.1 Role Summary')
    styled_table(doc, ['Role', 'Can Create Courses?', 'Can Publish?', 'Can View Reports?', 'Can Configure?'], [
        ("eLearning Administrator", "✔ Yes", "✔ Yes", "✔ Yes", "✔ Yes"),
        ("eLearning Manager",       "✔ Yes", "✔ Yes", "✔ Yes", "✔ Limited"),
        ("eLearning Officer",       "✔ Yes", "✔ Yes", "✔ Own Only", "✘ No"),
        ("Employee (Learner)",      "✘ No",  "✘ No",  "✘ No",       "✘ No"),
    ])

    add_h2(doc, '13.2 Assigning an eLearning Role')
    step(doc, 1, "Navigate to Settings → Users & Companies → Users.")
    step(doc, 2, "Open the user record.")
    step(doc, 3, "Scroll to the 'eLearning' section under permissions.")
    step(doc, 4, "Select the appropriate access level: User (Officer), Manager, or Administrator.")
    step(doc, 5, "Save. The user must log out and log back in for the change to take effect.")

    note(doc, "Employee/Learner access is the default for any user who does not have an explicit eLearning role. They can browse and enroll in published courses via the portal but cannot manage content.")
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
# CHAPTER 14 — BEST PRACTICES
# ─────────────────────────────────────────────────────────────
def ch14_best_practices(doc):
    add_h1(doc, '14. Best Practices & Tips')

    add_h2(doc, '14.1 Course Design Principles')
    for tip_text in [
        "Keep each slide/lesson focused on a single topic — 5 to 15 minutes of content per slide.",
        "Mix content types: start with a short video, follow with a document, end with a quiz.",
        "Use section dividers to chunk content into logical modules (e.g., Module 1, Module 2).",
        "Write descriptive course and slide titles — learners should know exactly what they will learn.",
        "Include a course description explaining prerequisites, learning outcomes, and expected time.",
        "Always include at least one quiz question per section to reinforce key learning points.",
    ]:
        bullet(doc, tip_text)

    add_h2(doc, '14.2 Video Best Practices')
    styled_table(doc, ['Topic', 'Recommendation'], [
        ("Format",         "MP4 with H.264 video codec and AAC audio codec"),
        ("Resolution",     "720p (1280×720) for most content; 1080p for detailed screen recordings"),
        ("Bitrate",        "1–2 Mbps for 720p; 4–8 Mbps for 1080p"),
        ("Max Duration",   "Under 15 minutes per video for best engagement; split longer content"),
        ("Max File Size",  "Under 500 MB preferred; up to 1.2 GB supported"),
        ("Audio",          "Clear narration with minimal background noise; use a quality microphone"),
        ("Subtitles",      "Add Arabic/Amharic captions for Ahadu-specific terminology where possible"),
    ])

    add_h2(doc, '14.3 Compliance Training Checklist')
    styled_table(doc, ['#', 'Step', 'Who'], [
        ("1",  "Identify mandatory compliance training topics for the year",               "HR / Compliance Officer"),
        ("2",  "Create course structure and sections before adding content",               "eLearning Officer"),
        ("3",  "Upload/link all video and document content",                               "eLearning Officer"),
        ("4",  "Add quiz questions for all compliance-critical slides",                    "eLearning Officer"),
        ("5",  "Preview the course as a learner and verify video playback",                "eLearning Manager"),
        ("6",  "Set Enroll Policy to 'On Invitation' for mandatory targeted training",     "eLearning Manager"),
        ("7",  "Publish the course",                                                       "eLearning Manager"),
        ("8",  "Send invitations to all target employees",                                 "eLearning Officer"),
        ("9",  "Monitor completion at Week 1, 2, and 3 after launch",                     "eLearning Manager"),
        ("10", "Send reminders to non-compliant employees",                                "eLearning Officer"),
        ("11", "Export completion report for audit record at deadline",                   "eLearning Manager"),
    ])

    add_h2(doc, '14.4 Troubleshooting Common Issues')
    styled_table(doc, ['Problem', 'Likely Cause', 'Solution'], [
        ("Video won't play",             "Wrong source type or corrupted file",         "Re-upload in MP4 format; check source type is set to 'local_file'"),
        ("Google Drive video shows error","Drive file is restricted access",             "Set Drive file sharing to 'Anyone with the link can view'"),
        ("Employee can't see the course","Course is Unpublished or wrong enroll policy","Publish the course; check and update enrollment policy"),
        ("Quiz answers not saving",      "Browser cache or session issue",              "Clear browser cache and retry; ask employee to use Chrome"),
        ("Completion % stuck at 0",      "Employee opened but didn't finish slides",    "Employee must complete all required slides including quizzes"),
        ("Tags not showing in filter",   "Tags not assigned to courses",               "Open each course and add the appropriate tags"),
        ("'Add Content' popup not working","Browser extension conflict or JS error",    "Disable browser extensions, especially ad-blockers, and retry"),
    ])
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
# CHAPTER 15 — GLOSSARY
# ─────────────────────────────────────────────────────────────
def ch15_glossary(doc):
    add_h1(doc, '15. Glossary')

    styled_table(doc, ['Term', 'Definition'], [
        ("Course / Channel",        "The top-level container grouping related lessons and training materials together"),
        ("Slide / Content",         "An individual piece of learning material: video, document, infographic, article, or webpage"),
        ("Section",                 "An organizational divider within a course that groups related slides (e.g., 'Module 1')"),
        ("Slide Category",          "The format/type of a slide: Video, Document, Infographic, Article, or Web Page"),
        ("Source Type",             "For Video slides: how the video is provided (local upload, Google Drive URL, or external link)"),
        ("Tag",                     "A label applied to a course or slide used for categorization and filtering"),
        ("Tag Group",               "A collection of related tags (e.g., 'Compliance' tag group containing 'AML', 'KYC' tags)"),
        ("Enroll Policy",           "The rule that controls who can join a course: Public, On Invitation, or On Payment"),
        ("Attendee / Learner",      "An employee who has enrolled in a course"),
        ("Completion Rate",         "Percentage of enrolled learners who have fully completed the course"),
        ("Quiz",                    "A set of multiple-choice questions attached to a slide for knowledge verification"),
        ("Karma",                   "Gamification points earned by learners for completing course activities"),
        ("Certificate",             "A digital document issued to a learner who completes a course"),
        ("Published",               "A course or slide that is visible and accessible to learners on the portal"),
        ("Unpublished",             "A course or slide in draft state, visible only to administrators and managers"),
        ("LMS",                     "Learning Management System — the platform for creating and delivering training courses"),
        ("HTML5 Player",            "The built-in browser-based video player used for locally uploaded video files"),
        ("Google Drive Preview",    "An embedded iframe player that streams video directly from Google Drive"),
        ("MIME Type",               "A technical format identifier for a file (e.g., video/mp4 for MP4 files)"),
        ("website_slides",          "The standard Odoo eLearning/LMS module that Ahadu eLearning extends"),
    ])


# ─────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────
def create_manual(output_path):
    doc = Document()

    # Page setup (A4)
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin  = section.bottom_margin = Cm(2.2)

    # Base font
    normal = doc.styles['Normal']
    normal.font.name  = 'Arial'
    normal.font.size  = Pt(11)
    normal.font.color.rgb = RGBColor(45, 55, 72)

    add_header_footer(doc)
    add_cover(doc)
    add_toc(doc)

    ch1_introduction(doc)
    ch2_getting_started(doc)
    ch3_dashboard(doc)
    ch4_courses(doc)
    ch5_content(doc)
    ch6_video(doc)
    ch7_tags(doc)
    ch8_quizzes(doc)
    ch9_enrollment(doc)
    ch10_reporting(doc)
    ch11_configuration(doc)
    ch12_publishing(doc)
    ch13_roles(doc)
    ch14_best_practices(doc)
    ch15_glossary(doc)

    doc.save(output_path)
    print(f"✅  Manual saved → {output_path}")


if __name__ == '__main__':
    out = '/opt/odoo18/odoo18/custom_addons/ahadu_elearning/Ahadu_eLearning_User_Manual_v2.docx'
    create_manual(out)
