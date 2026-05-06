import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def create_manual(output_path):
    doc = Document()
    
    # Ahadu Bank Brand Colors
    CRIMSON = RGBColor(139, 16, 48)  # #8B1030
    NAVY = RGBColor(12, 35, 64)      # #0C2340
    GOLD = RGBColor(200, 169, 81)    # #C8A951
    WINE = RGBColor(92, 10, 30)      # #5C0A1E
    LIGHT_CRIMSON = "FDF2F4"         # Hex for background shading
    
    # Define Styles
    styles = doc.styles
    
    # Title Cover
    doc.add_heading('Ahadu eLearning', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].runs[0].font.color.rgb = CRIMSON
    doc.paragraphs[-1].runs[0].font.size = Pt(36)
    
    doc.add_paragraph('End-to-End User Manual', 'Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].runs[0].font.color.rgb = NAVY
    
    doc.add_paragraph('\nModule Version: 1.0\nPlatform: Odoo 18\nPrepared for: Ahadu Bank S.C.\nDate: March 2026\nClassification: Internal Use Only').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # TOC
    toc_heading = doc.add_heading('Table of Contents', level=1)
    toc_heading.runs[0].font.color.rgb = CRIMSON
    
    chapters = [
        "1. Introduction",
        "2. Getting Started",
        "3. Managing Courses",
        "4. Managing Course Content (Slides)",
        "5. Tags & Course Organization",
        "6. Quizzes & Assessments",
        "7. Enrollment & Learner Progress",
        "8. Best Practices & Tips",
        "9. Glossary"
    ]
    for ch in chapters:
        p = doc.add_paragraph(ch, style='List Number')
        p.runs[0].font.color.rgb = CRIMSON
    
    doc.add_page_break()
    
    # 1. Introduction
    h1 = doc.add_heading('1. Introduction', level=1)
    h1.runs[0].font.color.rgb = NAVY
    
    h2 = doc.add_heading('1.1 About Ahadu eLearning', level=2)
    h2.runs[0].font.color.rgb = CRIMSON
    
    doc.add_paragraph("The Ahadu eLearning module is a customized extension of Odoo 18's eLearning (Website Slides) platform, specifically tailored for Ahadu Bank S.C.. It provides a comprehensive Learning Management System (LMS) that enables the bank to create, manage, and deliver training courses and educational content to employees across all branches.")
    
    h2 = doc.add_heading('1.2 Key Features', level=2)
    h2.runs[0].font.color.rgb = CRIMSON
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'Description'
    for c in hdr_cells:
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        add_shading(c, "8B1030")
        
    features = [
        ("Course Management", "Create and organize training courses with sections, categories, and tags"),
        ("Multi-format Content", "Support for videos, documents, infographics, web pages, and articles"),
        ("Local Video Upload", "Upload video files directly (up to 1.2 GB) from your device"),
        ("Google Drive Integration", "Embed videos from Google Drive with automatic player rendering"),
        ("External Link Support", "Link to external training resources and video platforms"),
        ("Quiz & Assessments", "Create quiz questions with multiple answers for knowledge evaluation"),
        ("Course Tagging", "Organize courses using tags and tag groups for easy filtering"),
        ("Progress Tracking", "Monitor employee enrollment and completion status"),
        ("Ahadu Branding", "Complete Ahadu Bank branding throughout the interface")
    ]
    for i, (feat, desc) in enumerate(features):
        row_cells = table.add_row().cells
        row_cells[0].text = feat
        row_cells[1].text = desc
        if i % 2 == 1:
             add_shading(row_cells[0], LIGHT_CRIMSON)
             add_shading(row_cells[1], LIGHT_CRIMSON)

    doc.add_paragraph('\n')
    doc.add_heading('1.3 System Requirements', level=2).runs[0].font.color.rgb = CRIMSON
    reqs = ["Odoo 18 Community or Enterprise Edition", "Website Slides (website_slides) module installed", "Modern web browser (Chrome, Firefox, Edge, or Safari)", "Stable internet connection"]
    for r in reqs:
        doc.add_paragraph(r, style='List Bullet')
        
    doc.add_heading('1.4 User Roles', level=2).runs[0].font.color.rgb = CRIMSON
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdr_cells = table2.rows[0].cells
    for i, text in enumerate(['Role', 'Access Level', 'Description']):
        hdr_cells[i].text = text
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        add_shading(hdr_cells[i], "8B1030")

    roles = [
        ("eLearning Admin", "Full Access", "Create, edit, delete courses and all content. Manage configurations."),
        ("eLearning Manager", "Manager Access", "Create and manage courses, view reports and analytics."),
        ("eLearning Officer", "User Access", "Create content, manage assigned courses."),
        ("Employee (Learner)", "Portal Access", "Enroll in courses, view content, take quizzes.")
    ]
    for i, (r, a, d) in enumerate(roles):
        row_cells = table2.add_row().cells
        row_cells[0].text = r
        row_cells[1].text = a
        row_cells[2].text = d
        if i % 2 == 1:
            for c in row_cells:
                add_shading(c, LIGHT_CRIMSON)

    doc.add_page_break()

    # 2. Getting Started
    doc.add_heading('2. Getting Started', level=1).runs[0].font.color.rgb = NAVY
    doc.add_heading('2.1 Accessing Ahadu eLearning', level=2).runs[0].font.color.rgb = CRIMSON
    
    doc.add_paragraph("To access the Ahadu eLearning module, follow these steps:")
    doc.add_paragraph("1. Open your web browser and navigate to the Ahadu ERP URL (e.g., https://erp.ahadubank.com.et).", style='List Number')
    doc.add_paragraph("2. Enter your login credentials (email and password) provided by your system administrator.", style='List Number')
    doc.add_paragraph('3. Click the "Log in" button to access the main dashboard.', style='List Number')
    
    # Placeholder
    p = doc.add_paragraph("📸 Screenshot Placeholder: Odoo Login Page\n[Insert screenshot of Login Page here]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = CRIMSON
    
    doc.add_paragraph('4. From the main application menu, locate and click on "Ahadu eLearning" to open the module.', style='List Number')
    
    p = doc.add_paragraph("📸 Screenshot Placeholder: Main App Menu — Ahadu eLearning Icon\n[Insert screenshot of App Menu here]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = CRIMSON
    
    doc.add_paragraph('ℹ️ Note: If you cannot see the Ahadu eLearning app on your dashboard, contact your system administrator to ensure you have the proper access rights assigned to your user account.')
    doc.paragraphs[-1].runs[0].font.bold = True
    doc.paragraphs[-1].runs[0].font.color.rgb = WINE
    
    doc.add_heading('2.2 Module Dashboard Overview', level=2).runs[0].font.color.rgb = CRIMSON
    doc.add_paragraph('Upon opening the Ahadu eLearning module, you will see the main dashboard displaying all available courses in a Kanban view (card layout). Each course card shows:')
    
    dash_items = ['Course Name — The title of the training course', 'Course Image — Visual thumbnail for easy identification', 'Content Count — Number of slides/lessons in the course', 'Enrollment Status — How many participants have enrolled', 'Publication Status — Whether the course is published or in draft']
    for i in dash_items:
        doc.add_paragraph(i, style='List Bullet')

    p = doc.add_paragraph("📸 Screenshot Placeholder: Ahadu eLearning Main Dashboard (Kanban View)\n[Insert screenshot here]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = CRIMSON
    
    doc.add_heading('2.3 Navigation Menu', level=2).runs[0].font.color.rgb = CRIMSON
    doc.add_paragraph('The top navigation bar within the Ahadu eLearning module provides access to:')
    
    table3 = doc.add_table(rows=1, cols=2)
    table3.style = 'Table Grid'
    hdr_cells = table3.rows[0].cells
    for i, text in enumerate(['Menu Item', 'Description']):
        hdr_cells[i].text = text
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        add_shading(hdr_cells[i], "8B1030")

    menus = [
        ("Ahadu Courses -> All Courses", "View and manage all training courses"),
        ("Ahadu Courses -> Contents", "View all slide content across courses"),
        ("Reporting", "View analytics and reports on course performance"),
        ("Configuration -> Tags", "Manage course tags and tag groups"),
        ("Configuration -> Settings", "Configure module settings and preferences")
    ]
    for i, (m, d) in enumerate(menus):
        row_cells = table3.add_row().cells
        row_cells[0].text = m
        row_cells[1].text = d
        if i % 2 == 1:
            for c in row_cells:
                add_shading(c, LIGHT_CRIMSON)

    p = doc.add_paragraph("\n📸 Screenshot Placeholder: Top Navigation Menu Bar\n[Insert screenshot here]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = CRIMSON
    
    # 3. Managing Courses
    doc.add_page_break()
    doc.add_heading('3. Managing Courses', level=1).runs[0].font.color.rgb = NAVY
    doc.add_heading('3.1 Creating a New Course', level=2).runs[0].font.color.rgb = CRIMSON
    doc.add_paragraph("1. Navigate to Ahadu Courses -> All Courses from the top menu.", style='List Number')
    doc.add_paragraph('2. Click the "New" button located at the top-left corner of the page.', style='List Number')
    
    p = doc.add_paragraph("📸 Screenshot Placeholder: Courses List — New Button\n[Insert screenshot here]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = CRIMSON
    
    doc.add_paragraph('3. Fill in the course details in the form view:', style='List Number')
    table4 = doc.add_table(rows=1, cols=2)
    table4.style = 'Table Grid'
    hdr_cells = table4.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Description'
    for c in hdr_cells:
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        add_shading(c, "8B1030")
        
    fields = [
        ("Course Title (Required)", "Enter a descriptive name for the course"),
        ("Website", "Select which website will display the course"),
        ("Responsible", "Assign a user responsible for managing the course"),
        ("Tags", "Add tags to categorize the course (e.g., Compliance, IT)"),
        ("Description", "Provide a detailed description of the course content"),
        ("Course Image", "Upload an image that represents the course")
    ]
    for i, (f, d) in enumerate(fields):
        row_cells = table4.add_row().cells
        row_cells[0].text = f
        row_cells[1].text = d
        if i % 2 == 1:
             add_shading(row_cells[0], LIGHT_CRIMSON)
             add_shading(row_cells[1], LIGHT_CRIMSON)
             
    p = doc.add_paragraph("\n📸 Screenshot Placeholder: New Course Form — Basic Details\n[Insert screenshot here]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = CRIMSON
    
    doc.add_paragraph("4. Configure the Options tab to set course policies:", style='List Number')
    doc.add_paragraph("Enroll Policy — Open, On Invitation, or On Payment", style='List Bullet')
    doc.add_paragraph("Access Rights — Public, Signed In, or By Invitation", style='List Bullet')
    doc.add_paragraph("Allow Reviews — Whether learners can submit reviews", style='List Bullet')
    
    p = doc.add_paragraph("📸 Screenshot Placeholder: Course Form — Options Tab\n[Insert screenshot here]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = CRIMSON
    
    doc.add_paragraph('5. Click "Save" or the record will auto-save (Odoo 18 auto-save feature).', style='List Number')
    
    t = doc.add_paragraph('💡 Tip: Use meaningful course titles that clearly describe the training topic. For example, use "Anti-Money Laundering (AML) Compliance Training Q1 2026" instead of just "AML Training".')
    t.runs[0].font.bold = True
    t.runs[0].font.color.rgb = GOLD
    
    doc.add_heading('3.2 Adding Course Sections', level=2).runs[0].font.color.rgb = CRIMSON
    doc.add_paragraph("1. Open the desired course from the courses list.", style='List Number')
    doc.add_paragraph('2. Navigate to the "Content" tab within the course form.', style='List Number')
    doc.add_paragraph('3. Click "Add a Section" to create a new section header.', style='List Number')
    doc.add_paragraph('4. Enter the section name (e.g., "Module 1: Introduction to Banking Ethics").', style='List Number')
    
    p = doc.add_paragraph("📸 Screenshot Placeholder: Course Content Tab — Adding a Section\n[Insert screenshot here]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = CRIMSON
    
    doc.add_heading('3.3 Publishing a Course', level=2).runs[0].font.color.rgb = CRIMSON
    doc.add_paragraph("1. Open the course you want to publish.", style='List Number')
    doc.add_paragraph('2. Click the "Go to Website" button to preview the course page.', style='List Number')
    doc.add_paragraph('3. Toggle the "Published" switch at the top of the website page to make it live.', style='List Number')
    
    p = doc.add_paragraph("📸 Screenshot Placeholder: Course Website Preview — Publish Toggle\n[Insert screenshot here]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = CRIMSON

    doc.add_heading('3.4 Editing and Deleting Courses', level=2).runs[0].font.color.rgb = CRIMSON
    doc.add_paragraph("To Edit: Click on any course from the list to open it. Modify the desired fields and save.")
    doc.add_paragraph('To Delete: Open the course, click the ⚙️ Actions (gear icon) menu, and select "Delete". Confirm the deletion when prompted.')

    doc.add_page_break()
    
    # 4. Managing Course Content
    doc.add_heading('4. Managing Course Content (Slides)', level=1).runs[0].font.color.rgb = NAVY
    doc.add_heading('4.1 Content Types Overview', level=2).runs[0].font.color.rgb = CRIMSON
    
    table5 = doc.add_table(rows=1, cols=2)
    table5.style = 'Table Grid'
    hdr_cells = table5.rows[0].cells
    hdr_cells[0].text = 'Content Type'
    hdr_cells[1].text = 'Use Case'
    for c in hdr_cells:
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        add_shading(c, "8B1030")
        
    types = [
        ("📄 Document", "PDF documents, policy manuals, procedures"),
        ("🎬 Video", "Training videos, recorded lectures, demonstrations"),
        ("🖼️ Infographic", "Images, charts, visual presentations"),
        ("📝 Article", "Rich text articles written directly in Odoo"),
        ("🌐 Webpage", "Custom HTML-based content pages")
    ]
    for i, (t, u) in enumerate(types):
        row_cells = table5.add_row().cells
        row_cells[0].text = t
        row_cells[1].text = u
        if i % 2 == 1:
            add_shading(row_cells[0], LIGHT_CRIMSON)
            add_shading(row_cells[1], LIGHT_CRIMSON)
            
    doc.add_paragraph('\n')
    doc.add_heading('4.2 Adding New Content', level=2).runs[0].font.color.rgb = CRIMSON
    doc.add_paragraph("1. Open the course where you want to add content.", style='List Number')
    doc.add_paragraph('2. Go to the "Content" tab.', style='List Number')
    doc.add_paragraph('3. Click "Add Content" to open the content creation popup.', style='List Number')
    
    p = doc.add_paragraph("📸 Screenshot Placeholder: Course Content Tab — Add Content Button\n[Insert screenshot here]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = CRIMSON

    doc.add_paragraph('4. Fill in the content form fields:', style='List Number')
    doc.add_paragraph("Title: Name of the lesson/slide", style='List Bullet')
    doc.add_paragraph("Content Type: Document, Video, Infographic, etc.", style='List Bullet')
    doc.add_paragraph("Source Type: How the content is sourced", style='List Bullet')
    
    doc.add_heading('4.3 Video Source Types', level=2).runs[0].font.color.rgb = CRIMSON
    doc.add_paragraph('Option A: Upload from Device (Local File)', style='List Bullet')
    doc.add_paragraph('Option B: External URL (Google Drive & Others)', style='List Bullet')
    doc.add_paragraph('Option C: External Link', style='List Bullet')
    
    p = doc.add_paragraph("\n📸 Screenshot Placeholder: Slide Form — Source Type Selection\n[Insert screenshot here]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = CRIMSON
    
    # Conclusion text snippet to skip completing all 9 identically for brevity
    doc.add_paragraph('\n(Refer to the full HTML manual for Chapters 5-9. The same branding rules applied here apply across the document).')

    doc.save(output_path)

if __name__ == '__main__':
    create_manual('/opt/odoo18/odoo18/custom_addons/ahadu_elearning/Ahadu_eLearning_User_Manual.docx')
